from __future__ import annotations

from pathlib import Path


class TextExtractionError(RuntimeError):
    pass


class TextExtractor:
    IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}

    def extract(self, file_path: str | Path) -> tuple[str, list[str]]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file does not exist: {path}")

        suffix = path.suffix.lower()
        warnings: list[str] = []

        if suffix in {".txt", ".md"}:
            return self._extract_text_file(path), warnings
        if suffix == ".pdf":
            return self._extract_pdf(path), warnings
        if suffix == ".docx":
            return self._extract_docx(path), warnings
        if suffix in self.IMAGE_SUFFIXES:
            return self._extract_image(path), warnings

        raise TextExtractionError(f"Unsupported resume file type: {suffix or '<none>'}")

    def _extract_text_file(self, path: Path) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_text(errors="ignore")

    def _extract_pdf(self, path: Path) -> str:
        try:
            import pdfplumber
        except ImportError as exc:
            raise TextExtractionError(
                "PDF parsing requires pdfplumber. Install with: pip install -e \".[full]\""
            ) from exc

        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return "\n".join(pages).strip()

    def _extract_docx(self, path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise TextExtractionError(
                "DOCX parsing requires python-docx. Install with: pip install -e \".[full]\""
            ) from exc

        document = Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(paragraphs + table_cells).strip()

    def _extract_image(self, path: Path) -> str:
        try:
            from PIL import Image
            import pytesseract
        except ImportError as exc:
            raise TextExtractionError(
                "Image OCR requires pillow and pytesseract. Install with: pip install -e \".[full]\""
            ) from exc

        with Image.open(path) as image:
            return pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
