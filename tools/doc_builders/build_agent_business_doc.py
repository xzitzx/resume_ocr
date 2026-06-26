from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:/Users/zhangxu/Desktop/workspace/jianli_ocr")
OUT = ROOT / "output" / "documents" / "霖珑云科淮安分部Agent业务应用方案.docx"
MINDMAP = Path(r"C:/Users/zhangxu/Desktop/workspace/微信图片_20260626092443.png")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 42)
MUTED = RGBColor(90, 96, 106)
LIGHT = "F2F4F7"
LIGHT_BLUE = "E8EEF5"


def set_east_asia(run, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_para(p, before=0, after=6, line=1.1):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def add_run(p, text, size=11, bold=False, color=INK):
    r = p.add_run(text)
    set_east_asia(r)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    return r


def add_p(doc, text="", size=11, bold=False, color=INK, after=6, before=0, align=None):
    p = doc.add_paragraph()
    style_para(p, before=before, after=after)
    if align is not None:
        p.alignment = align
    add_run(p, text, size=size, bold=bold, color=color)
    return p


def add_h(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.clear()
    if level == 1:
        size, color, before, after = 16, BLUE, 16, 8
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, DARK_BLUE, 8, 4
    style_para(p, before=before, after=after, line=1.1)
    add_run(p, text, size=size, bold=True, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style_para(p, after=4, line=1.167)
        add_run(p, item, size=10.5)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        style_para(p, after=4, line=1.167)
        add_run(p, item, size=10.5)


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_para(p, after=3)
    add_run(p, title + "：", size=10.5, bold=True, color=DARK_BLUE)
    add_run(p, body, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_para(p, after=0, line=1.1)
        add_run(p, h, size=9.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(text)) > 12 else WD_ALIGN_PARAGRAPH.CENTER
            style_para(p, after=0, line=1.1)
            add_run(p, str(text), size=9.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name in ["List Bullet", "List Number"]:
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.167
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(header, "霖珑云科淮安分部 Agent 业务应用方案", size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(footer, "第 ", size=9, color=MUTED)
    add_page_number(footer)
    add_run(footer, " 页", size=9, color=MUTED)
    return doc


def build():
    doc = setup_doc()

    add_p(doc, "内部业务方案", size=11, bold=True, color=BLUE, after=8)
    p = add_p(doc, "Agent 在霖珑云科业务中的应用落地方案", size=24, bold=True, color=INK, after=6)
    p.paragraph_format.space_before = Pt(18)
    add_p(doc, "面向淮安分部、金蜻蜓即配骑手服务、薪社汇、乐接活、小霖数字员工及 AI 创新联合体场景", size=12.5, color=MUTED, after=18)

    rows = [
        ("使用对象", "淮安分部 Agent 开发工程师、业务负责人、产品经理、交付负责人、风控合规团队"),
        ("目标", "形成可执行的 Agent 应用蓝图，帮助业务提效、风控前置、经验沉淀与产品化复制"),
        ("资料依据", "霖珑云科公司介绍 2024、霖珑云科淮安材料、数智创新数字化 2026 业务脑图"),
        ("生成日期", "2026-06-26"),
    ]
    add_matrix(doc, ["项目", "说明"], rows, [1.45, 5.05])

    add_callout(
        doc,
        "核心判断",
        "霖珑云科做 Agent 的最大机会，不是做一个通用聊天机器人，而是把 Agent 嵌入薪税社保、灵活用工、即配骑手服务、属地交付、合规风控和政府产业合作流程中，形成“懂业务、能办事、可审计、可复制”的业务执行型智能体体系。",
    )

    add_h(doc, "目录", 1)
    toc_items = [
        "一、业务背景与机会判断",
        "二、Agent 在公司业务中的总体定位",
        "三、淮安分部优先落地场景",
        "四、重点 Agent 方案设计",
        "五、系统架构与技术实现建议",
        "六、数据、知识库与工具体系",
        "七、风控、安全与合规治理",
        "八、产品化与商业化路径",
        "九、90 天行动计划",
        "十、开发工程师能力建设建议",
        "附录：场景清单、指标体系、业务脑图",
    ]
    add_numbers(doc, toc_items)

    doc.add_page_break()

    add_h(doc, "一、业务背景与机会判断", 1)
    add_h(doc, "1.1 公司业务的本质", 2)
    add_p(doc, "从公司材料看，霖珑云科不是单一软件公司，也不是传统人力外包公司，而是一家“技术 + 服务”双驱动的人力资源科技企业。其业务一端连接企业客户，一端连接员工、灵工、骑手、属地服务人员、政府与生态伙伴，核心价值在于把复杂的人力资源服务流程数字化、平台化、合规化。")
    add_p(doc, "集团已有薪社汇、乐接活、微办公、超级 HR、安心工资条、AI 数字员工“小霖”等产品和品牌，覆盖标准用工、灵活用工、出海用工、协同办公、数字人事、电子薪资单等场景。淮安公司则围绕金蜻蜓即配骑手服务、新业态合作基地、人力资源产业园、AI 创新联合体等方向，具备非常清晰的业务抓手。")
    add_callout(doc, "机会结论", "Agent 最适合介入“规则复杂、资料繁多、跨系统流转、需要经验判断、但最终仍可被流程化沉淀”的场景。霖珑云科的薪税社保、灵活用工、即配服务、保险理赔、政策咨询和属地交付，恰好都属于这类场景。")

    add_h(doc, "1.2 为什么不是普通 AI 问答，而是业务型 Agent", 2)
    rows = [
        ("普通问答机器人", "回答问题为主，通常不理解业务上下文，不负责流程结果。", "适合 FAQ、知识查询、售前咨询。"),
        ("业务 Copilot", "辅助业务人员完成判断、检索、整理、生成材料。", "适合客服、交付、方案顾问、运营人员。"),
        ("业务执行型 Agent", "可调用系统、读取数据、触发流程、生成工单、进行异常检测，并把结果留痕。", "适合社保办理、理赔材料审核、风控预警、发薪报税检查。"),
        ("多 Agent 协同网络", "多个专业 Agent 分工协作，由调度 Agent 编排任务。", "适合金蜻蜓骑手全周期服务、政产学研平台、全国属地交付。"),
    ]
    add_matrix(doc, ["类型", "特点", "适配霖珑场景"], rows, [1.35, 2.55, 2.6])
    add_p(doc, "公司当前已经具备 SaaS 产品、业务数据、属地服务经验、政策知识和客户场景，下一步重点应从“AI 能回答什么”转向“AI 能帮助业务完成什么”。这也是 Agent 相比传统机器人最大的差异。")

    add_h(doc, "1.3 淮安分部的特殊价值", 2)
    add_bullets(doc, [
        "淮安金蜻蜓聚焦即时配送骑手服务，业务链路完整，覆盖招聘、运单、保险投保、日常理赔、交安处理、新职伤保障、发薪报税等场景，天然适合做流程型 Agent。",
        "淮安材料中提出 AI 创新联合体、AI 社保智能体、AI 就业平台、数字仲裁、企业数字化转型等方向，说明本地不仅有内部提效需求，也有对外示范和政企合作空间。",
        "即配业务数据高频、异常多、规则强、规模增长快。Agent 如果能在这里跑通，就可以向乐接活、薪社汇、其他城市即配项目复制。",
        "淮安分部可以成为集团 Agent 产品化的“业务试验田”和“样板间”：先内部使用，再固化为产品能力，最后服务外部客户和政府合作项目。",
    ])

    add_h(doc, "二、Agent 在公司业务中的总体定位", 1)
    add_h(doc, "2.1 定位一句话", 2)
    add_callout(doc, "总体定位", "霖珑云科 Agent 应定位为“懂人力资源业务、懂属地政策、能调用业务系统、能沉淀专家经验、可审计可追责的业务执行型数字员工”。")
    add_h(doc, "2.2 五层能力模型", 2)
    rows = [
        ("感知层", "读取客户问题、骑手材料、合同、运单、保单、工单、政策文件、表格、图片和系统数据。", "OCR、文档解析、结构化抽取、多模态识别。"),
        ("知识层", "理解公司产品、政策法规、SOP、案例、城市规则、合同模板和风险准入标准。", "RAG、知识图谱、向量检索、规则库。"),
        ("推理层", "根据业务上下文判断合规性、材料完整性、风险等级、办理路径和下一步动作。", "大模型推理、规则引擎、专家模型、评分卡。"),
        ("执行层", "调用 SaaS 系统、工单系统、CRM、发薪报税、保险理赔、消息通知、审批流等工具。", "Function Calling、Workflow、API 编排、RPA。"),
        ("治理层", "记录依据、来源、版本、操作人、审批人、模型输出和人工修正结果。", "审计日志、权限控制、灰度发布、质检闭环。"),
    ]
    add_matrix(doc, ["层级", "业务含义", "技术抓手"], rows, [1.0, 3.1, 2.4])

    add_h(doc, "2.3 三条设计原则", 2)
    add_numbers(doc, [
        "先做业务闭环，再追求智能炫技。每个 Agent 都必须对应一个明确业务结果，例如减少人工审核时长、降低漏审率、提升一次性材料通过率、缩短工单处理周期。",
        "高风险事项人审兜底。涉及劳动关系认定、税务合规、发薪、报税、保险拒赔、法律责任等事项，Agent 只能给建议、证据和风险等级，最终决策必须由授权人员确认。",
        "知识、规则、工具一体化。Agent 不能只接大模型，必须同时连接知识库、规则引擎、业务数据库和系统工具，否则只能聊天，不能办事。",
    ])

    add_h(doc, "三、淮安分部优先落地场景", 1)
    add_p(doc, "结合淮安材料和业务脑图，建议按照“金蜻蜓即配业务优先、内部客服交付次之、政企示范同步储备”的顺序推进。")
    rows = [
        ("P0", "骑手理赔材料审核 Agent", "事故、理赔、新职伤材料多，人工重复检查多。", "材料完整率、审核时长、一次通过率", "强烈建议优先做 MVP"),
        ("P0", "社保薪税政策助手 Agent", "全国政策差异大，客服与交付咨询频繁。", "答复准确率、人工升级率、平均响应时长", "适合与薪社汇、小霖结合"),
        ("P0", "灵活用工准入风控 Agent", "场景合规是乐接活和即配业务底线。", "高风险识别率、准入误判率、整改闭环率", "需要规则与专家审核结合"),
        ("P1", "骑手招聘与准入 Agent", "招聘规模化后资料收集、资格初筛和沟通成本高。", "入职转化率、资料补齐时长", "可连接新媒体招聘"),
        ("P1", "运单异常检测 Agent", "异常运单会影响发薪、保险、合规与客户结算。", "异常召回率、误报率、处理时长", "需要数据接口支持"),
        ("P1", "客服交付 Copilot", "客服需要查客户、查政策、查流程、写纪要。", "工单关闭时长、满意度、质检得分", "见效快，适合内部推广"),
        ("P2", "企业方案生成 Agent", "售前方案高度依赖专家经验。", "方案产出时长、采纳率、商机转化", "适合沉淀顾问方法论"),
        ("P2", "AI 就业匹配 Agent", "面向政府、高校、新业态合作基地。", "匹配成功率、培训就业转化", "适合政企示范项目"),
    ]
    add_matrix(doc, ["优先级", "场景", "业务痛点", "核心指标", "建议"], rows, [0.7, 1.45, 1.8, 1.35, 1.2])

    add_h(doc, "四、重点 Agent 方案设计", 1)
    add_h(doc, "4.1 金蜻蜓骑手全周期服务 Agent 群", 2)
    add_p(doc, "金蜻蜓业务覆盖骑手“招、保、管、薪”全周期服务，是淮安分部最适合打造 Agent 样板的业务线。建议不是做一个大而全 Agent，而是按照业务环节拆成多个专业 Agent，由一个调度 Agent 统一编排。")
    rows = [
        ("调度 Agent", "理解用户意图，判断应调用招聘、运单、保险、理赔、发薪或风控 Agent。", "业务人员一句话发起任务，系统自动拆解流程。"),
        ("招聘 Agent", "筛选候选人资料、判断准入条件、生成沟通话术、提醒补充材料。", "减少招聘专员重复沟通，提高入职资料完整度。"),
        ("运单 Agent", "识别重复单、异常轨迹、金额异常、时间异常、骑手与订单不匹配。", "提前发现影响结算和风控的问题。"),
        ("保险 Agent", "核对投保状态、保障范围、保单有效期、骑手身份信息一致性。", "降低漏保、错保和责任不清风险。"),
        ("理赔 Agent", "读取事故描述、图片、证明材料，判断材料缺失和理赔路径。", "提升理赔材料一次通过率。"),
        ("新职伤 Agent", "根据事故时间、地点、任务关系和政策规则判断是否建议申报。", "减少人工初判压力，保留判断依据。"),
        ("发薪报税 Agent", "检查运单、佣金、发票、完税、个税和支付流水一致性。", "支撑四流合一与税务合规。"),
        ("风险 Agent", "汇总个人、订单、事故、投诉、结算异常，输出风险等级和处置建议。", "为运营主管提供风险看板。"),
    ]
    add_matrix(doc, ["Agent", "职责", "业务价值"], rows, [1.15, 2.85, 2.5])

    add_h(doc, "4.2 骑手理赔材料审核 Agent 详细流程", 2)
    add_numbers(doc, [
        "输入：骑手身份信息、事故时间地点、运单记录、保单信息、事故照片、诊断证明、费用票据、交警责任认定、客服沟通记录。",
        "识别：OCR 读取图片和票据，抽取姓名、身份证、日期、金额、医院、事故类型、责任比例、票据编号等字段。",
        "校验：比对骑手身份、保单有效期、事故时间与运单时间、票据时间与诊疗时间、金额区间和材料完整性。",
        "判断：输出案件类型、材料缺失项、疑似风险点、建议理赔路径、是否需要人工复核。",
        "生成：自动生成材料清单、补充材料短信、内部审核摘要、客户侧说明和工单备注。",
        "留痕：保存模型输出、引用依据、人工修改记录、最终处理结果，用于后续质检和模型优化。",
    ])
    rows = [
        ("低风险", "材料齐全，字段一致，事故在保障范围内，金额正常。", "Agent 生成审核摘要，人工快速确认。"),
        ("中风险", "材料缺失、字段小幅不一致、票据模糊、事故描述不完整。", "Agent 发起补件，人工复核关键字段。"),
        ("高风险", "疑似非任务期间事故、保单失效、身份不一致、金额异常、重复报案。", "强制转交专员或法务/风控，不允许自动通过。"),
    ]
    add_matrix(doc, ["风险等级", "判断特征", "处理方式"], rows, [1.0, 3.0, 2.5])

    add_h(doc, "4.3 社保薪税政策 Agent", 2)
    add_p(doc, "薪社汇和属地交付业务的难点在于政策地域差异、客户问题复杂、更新频率高。建议建设“政策 RAG + 城市规则库 + 办理 SOP + 人工确认”的社保薪税政策 Agent。")
    add_bullets(doc, [
        "面向客服：快速回答客户咨询，生成标准话术和办理清单。",
        "面向 HR 客户：提供城市政策、增减员、补缴、基数调整、公积金、个税等查询。",
        "面向交付人员：提示办理窗口、材料要求、截止日期、异常处理流程。",
        "面向管理者：汇总高频政策问题、政策变化影响客户数量、待更新 SOP。",
    ])
    add_callout(doc, "关键要求", "政策 Agent 必须显示依据来源、适用城市、政策时间、置信度和人工确认状态。没有来源或过期政策不能直接对外输出。")

    add_h(doc, "4.4 灵活用工准入与合规风控 Agent", 2)
    add_p(doc, "乐接活和即配业务都涉及非劳动关系、平台众包、业务外包、自然人经营行为、个体户经营行为等复杂模式。Agent 应前置到商机和项目准入阶段，先判断能不能做，再讨论怎么做。")
    rows = [
        ("真实业务场景", "是否存在真实任务、真实交付、真实验收。", "无真实业务或仅为发薪通道。"),
        ("管理强度", "是否强排班、强考勤、强纪律、强绩效。", "接近劳动关系管理方式。"),
        ("人员身份", "是否存在禁入人群、在职员工转灵工、特殊年龄或资质要求。", "劳动关系混同或身份不合规。"),
        ("四流一致", "合同流、业务流、资金流、票据流是否闭环。", "发票、资金、任务、合同无法相互印证。"),
        ("税务路径", "自然人代征、小额零星、临时税务登记、个体工商户是否匹配。", "税务处理方式与实际场景不符。"),
        ("纠纷风险", "是否存在伤害、投诉、仲裁、舆情、客户违约历史。", "已有高频纠纷或重大舆情风险。"),
    ]
    add_matrix(doc, ["检查维度", "Agent 判断内容", "高风险信号"], rows, [1.2, 2.7, 2.6])

    add_h(doc, "4.5 客服交付 Copilot", 2)
    add_p(doc, "客服和属地交付是公司服务体验的关键。Copilot 不替代客服，而是把客服从查资料、写纪要、重复解释中释放出来。")
    add_bullets(doc, [
        "客户进入会话时，自动拉取客户档案、项目类型、所在城市、历史工单、合同状态和最近异常。",
        "客服提问后，Agent 结合政策库、SOP、客户上下文生成建议回复，并标注是否可直接发送。",
        "会话结束后，自动生成工单摘要、待办事项、责任人、截止时间和风险标签。",
        "对超时、投诉、反复咨询、政策争议、疑似法务风险的工单自动提醒主管。",
    ])

    add_h(doc, "4.6 企业方案生成 Agent", 2)
    add_p(doc, "销售和售前面对客户时，经常需要根据行业、城市、人数、用工模式、成本目标和风险偏好设计方案。方案生成 Agent 可以把优秀顾问的经验产品化。")
    add_numbers(doc, [
        "输入客户画像：行业、城市、人数、岗位、用工波动、是否多分支、是否平台型、当前痛点。",
        "匹配业务模型：标准用工、非全日制、劳务派遣、岗位外包、灵活用工、即配骑手服务、出海用工等。",
        "生成方案结构：现状诊断、推荐模式、流程设计、系统能力、服务边界、合规风险、预计收益。",
        "输出材料：客户版方案、内部评审版、报价辅助、实施计划、风险提示清单。",
    ])

    add_h(doc, "五、系统架构与技术实现建议", 1)
    add_h(doc, "5.1 推荐总体架构", 2)
    rows = [
        ("用户入口", "微办公、小霖、企业微信/微信、客服系统、业务后台、移动端。", "统一身份、权限、会话入口。"),
        ("Agent 编排层", "意图识别、任务拆解、Agent 路由、状态机、流程编排、人工接管。", "建议自研轻量编排框架，重要流程可配置。"),
        ("模型层", "通用大模型、垂直小模型、OCR、多模态模型、Embedding 模型。", "按成本和风险分层调用。"),
        ("知识层", "政策库、SOP、案例库、合同模板、城市规则、产品说明、FAQ。", "RAG + 版本管理 + 审核发布。"),
        ("工具层", "CRM、工单、薪社汇、乐接活、保险、发薪、报税、消息、审批、文件系统。", "通过 API 或 RPA 封装为 Tool。"),
        ("数据层", "客户、骑手、运单、合同、保单、工单、发薪、发票、完税、日志。", "做统一主数据和字段字典。"),
        ("治理层", "权限、审计、质检、评测、灰度、日志、敏感信息保护。", "高风险动作必须留痕和审批。"),
    ]
    add_matrix(doc, ["层级", "内容", "建设重点"], rows, [1.0, 3.2, 2.3])

    add_h(doc, "5.2 技术选型建议", 2)
    add_bullets(doc, [
        "大模型：采用“通用大模型 + 专家规则 + 企业知识库”的组合，不建议完全依赖单一大模型。",
        "RAG：政策、SOP、合同、案例类知识必须结构化入库，支持来源追踪、版本号、有效期、城市标签和业务线标签。",
        "工作流：涉及办理、审批、补件、转人工的场景，用状态机或工作流引擎管理，不把流程藏在 Prompt 里。",
        "工具调用：每个业务系统能力封装为明确 Tool，例如查询保单、创建工单、生成补件通知、查询社保规则、校验运单。",
        "规则引擎：劳动关系、税务准入、四流合一、理赔规则等高风险判断应以规则引擎为主，大模型负责解释和辅助判断。",
        "评测体系：为每个 Agent 建测试集，覆盖常见问题、边界问题、政策更新、异常数据和高风险拒答场景。",
    ])

    add_h(doc, "5.3 Agent 开发模式", 2)
    add_p(doc, "建议采用“场景定义 -> 知识准备 -> 工具封装 -> Prompt/规则设计 -> 人工审核 -> 小范围试点 -> 指标评估 -> 产品化”的闭环方式。")
    rows = [
        ("需求定义", "明确用户、输入、输出、边界、成功指标。", "业务流程图、字段清单、风险清单。"),
        ("知识准备", "整理政策、SOP、案例、模板、FAQ。", "可检索知识库、版本记录。"),
        ("工具封装", "把系统查询、写入、通知、审批能力封装。", "Tool API 文档、权限说明。"),
        ("Agent 设计", "编写系统提示词、规则、状态机、异常处理。", "Agent 配置、测试用例。"),
        ("评测验证", "用历史案例和人工标注集测试。", "准确率、召回率、拒答率、人工满意度。"),
        ("试点上线", "选择一个团队、一个城市或一个客户试点。", "周报、问题清单、优化计划。"),
        ("规模推广", "沉淀为产品模块和运营 SOP。", "版本路线图、培训材料、质检规则。"),
    ]
    add_matrix(doc, ["阶段", "工作内容", "产出物"], rows, [1.0, 3.0, 2.5])

    add_h(doc, "六、数据、知识库与工具体系", 1)
    add_h(doc, "6.1 知识库分层", 2)
    rows = [
        ("集团通用知识", "公司介绍、产品矩阵、服务边界、品牌口径、通用 FAQ。", "市场、销售、客服、小霖。"),
        ("业务线知识", "薪社汇、乐接活、金蜻蜓、出海用工、微办公等产品 SOP。", "业务人员、产品经理、实施交付。"),
        ("政策法规知识", "社保、公积金、个税、灵活就业、新职伤、劳动关系、税务政策。", "客服、法务、风控、交付。"),
        ("城市属地知识", "各城市材料、窗口、办理周期、截止时间、特殊规则。", "属地服务团队。"),
        ("案例知识", "客户案例、理赔案例、仲裁案例、风控案例、异常处理案例。", "专家复盘、模型评测、培训。"),
        ("模板知识", "合同模板、通知模板、补件模板、方案模板、报告模板。", "自动生成材料。"),
    ]
    add_matrix(doc, ["知识层", "内容", "服务对象"], rows, [1.15, 3.4, 1.95])

    add_h(doc, "6.2 数据打通优先级", 2)
    add_numbers(doc, [
        "第一优先级：客户、骑手、工单、运单、保单、理赔、发薪报税等与金蜻蜓 P0 场景直接相关的数据。",
        "第二优先级：社保薪税政策、城市属地规则、客户合同、SOP、历史咨询和案例库。",
        "第三优先级：CRM 商机、销售方案、报价、企业画像、人才画像、就业培训数据。",
        "第四优先级：政府数据、产业报告、就业指数、产教融合数据，主要服务外部示范项目。",
    ])

    add_h(doc, "6.3 工具调用清单", 2)
    rows = [
        ("查询类", "查客户、查骑手、查运单、查保单、查工单、查政策、查合同、查发薪记录。"),
        ("生成类", "生成补件通知、审核摘要、客服话术、客户方案、日报周报、风险报告。"),
        ("校验类", "校验身份、保单有效期、材料完整性、四流一致、政策适用性、重复报案。"),
        ("流转类", "创建工单、转人工、发起审批、发送短信/微信、同步 CRM、更新处理状态。"),
        ("监控类", "异常预警、超时提醒、风险看板、质检抽样、模型输出审计。"),
    ]
    add_matrix(doc, ["工具类别", "典型能力"], rows, [1.2, 5.3])

    add_h(doc, "七、风控、安全与合规治理", 1)
    add_p(doc, "霖珑云科的业务涉及劳动关系、社保、税务、薪酬、保险、发票、个体工商户、骑手伤害、仲裁等高敏领域。Agent 的安全设计必须从第一天纳入系统，而不是上线后补救。")
    add_h(doc, "7.1 高风险动作边界", 2)
    rows = [
        ("可自动完成", "资料格式检查、知识检索、摘要生成、低风险提醒、工单草稿、补件清单。", "记录输出即可。"),
        ("需人工确认", "政策适用判断、理赔建议、发薪报税异常处理、客户方案、风险等级调整。", "业务负责人或专员确认。"),
        ("禁止自动决策", "劳动关系最终认定、税务合规承诺、拒赔结论、仲裁/诉讼意见、对外法律承诺。", "法务、风控或授权负责人决策。"),
    ]
    add_matrix(doc, ["动作等级", "示例", "治理要求"], rows, [1.2, 3.4, 1.9])

    add_h(doc, "7.2 审计留痕要求", 2)
    add_bullets(doc, [
        "每次模型输出必须记录：用户、时间、输入、引用知识、工具调用、模型版本、输出结果、置信度。",
        "每次人工修改必须记录：修改人、修改内容、修改原因、最终采用结果。",
        "每个政策知识条目必须记录：来源、发布日期、适用城市、有效期、审核人、版本号。",
        "每个高风险案件必须记录：Agent 判断、规则命中、人工复核、最终处理和复盘结果。",
    ])

    add_h(doc, "7.3 数据安全", 2)
    add_bullets(doc, [
        "个人敏感信息脱敏：身份证、手机号、银行卡、病历、事故照片等数据展示和训练都要脱敏或最小化使用。",
        "权限隔离：客服、交付、销售、风控、法务、管理者看到的数据范围应不同。",
        "模型供应商隔离：外部模型调用不得传输不必要的敏感字段，高敏场景优先使用私有化或脱敏后调用。",
        "训练数据治理：未经授权的客户合同、个人资料、理赔材料不得直接进入模型训练集。",
    ])

    add_h(doc, "八、产品化与商业化路径", 1)
    add_h(doc, "8.1 内部提效到外部产品", 2)
    add_p(doc, "Agent 能力最稳妥的商业化路径，是先在内部真实业务中验证，再把稳定能力封装为产品模块。")
    rows = [
        ("内部工具期", "面向客服、交付、运营、风控人员使用。", "验证准确性、效率提升和流程适配。", "内部 Copilot、后台插件。"),
        ("产品模块期", "嵌入薪社汇、乐接活、金蜻蜓、小霖、微办公。", "形成客户可感知能力。", "AI 政策助手、AI 理赔助手、AI 风控助手。"),
        ("行业方案期", "按行业打包为即配、物流、连锁门店、制造外包、新零售方案。", "提升销售转化和客单价。", "行业 Agent 套件。"),
        ("政企生态期", "服务产业园、就业平台、数字政府治理、产教融合。", "打造淮安示范项目。", "AI 就业平台、社保智能体、产业报告。"),
    ]
    add_matrix(doc, ["阶段", "对象", "目标", "产品形态"], rows, [1.0, 2.0, 1.8, 1.7])

    add_h(doc, "8.2 可包装的产品能力", 2)
    add_bullets(doc, [
        "金蜻蜓 AI 骑手服务中台：招聘、保险、理赔、新职伤、发薪报税、风险看板。",
        "薪社汇 AI 政策助手：全国社保薪税政策查询、办理清单、材料校验、异常预警。",
        "乐接活 AI 合规引擎：场景准入、四流一致、灵工风险、税务路径建议。",
        "小霖数字员工升级版：从问答助手升级为能调用业务系统的任务型数字员工。",
        "AI 就业与培训平台：人才画像、岗位画像、智能匹配、培训路径、就业跟踪。",
    ])

    add_h(doc, "九、90 天行动计划", 1)
    add_p(doc, "建议淮安分部用 90 天跑通第一个可验证 MVP，不追求一次性覆盖所有业务。首选方向：骑手理赔材料审核 Agent 或社保薪税政策助手 Agent。")
    rows = [
        ("第 1-2 周", "确定 MVP 场景和指标", "选择一个高频流程，明确输入输出、人工流程、风险边界。", "需求说明、流程图、指标口径。"),
        ("第 3-4 周", "准备数据和知识", "整理 50-100 个历史案例、SOP、材料模板、规则清单。", "知识库、测试集、字段字典。"),
        ("第 5-6 周", "完成原型", "搭建 RAG、规则校验、基础工具调用和审核摘要生成。", "可演示原型。"),
        ("第 7-8 周", "内部试用", "选择 3-5 名业务人员试用，记录问题和人工修正。", "试用报告、优化清单。"),
        ("第 9-10 周", "强化风控和留痕", "加入权限、审计、置信度、人工确认和异常转派。", "可控上线版本。"),
        ("第 11-12 周", "小范围上线", "在一个业务组或一个项目上线，按周看指标。", "上线复盘、下一阶段路线图。"),
    ]
    add_matrix(doc, ["周期", "主题", "主要工作", "交付物"], rows, [0.9, 1.4, 2.8, 1.4])

    add_h(doc, "9.1 MVP 成功指标", 2)
    add_bullets(doc, [
        "效率指标：单件处理时长下降 30% 以上。",
        "质量指标：材料缺失识别准确率达到 85% 以上，高风险漏判率接近 0。",
        "体验指标：业务人员主观满意度达到 4/5 以上。",
        "复用指标：至少沉淀 1 套知识库、1 套规则清单、1 套评测集、1 套标准流程。",
        "管理指标：所有高风险输出均有人工确认和审计记录。",
    ])

    add_h(doc, "十、开发工程师能力建设建议", 1)
    add_p(doc, "作为淮安分部 Agent 开发工程师，最核心的竞争力不是只会调模型接口，而是能把业务流程、系统工具、知识库和风控规则组合成真正可用的业务产品。")
    rows = [
        ("业务理解", "懂薪税社保、灵活用工、即配骑手、保险理赔、属地交付基本流程。", "每周跟客服/运营/理赔坐席旁听，形成场景笔记。"),
        ("Agent 工程", "掌握 RAG、工具调用、状态机、多 Agent 编排、Prompt 评测。", "做一个可复用 Agent 模板工程。"),
        ("数据工程", "能处理文档、表格、图片、业务数据库和日志。", "建立字段字典和数据质量检查脚本。"),
        ("风控意识", "理解哪些可自动化，哪些必须人审。", "和法务/风控共同维护红线规则。"),
        ("产品思维", "能把内部需求转化为可复制产品能力。", "每个 Agent 都配指标、流程、权限和培训材料。"),
    ]
    add_matrix(doc, ["能力", "要求", "建议动作"], rows, [1.1, 2.8, 2.6])

    add_h(doc, "附录 A：可落地 Agent 场景清单", 1)
    rows = [
        ("薪社汇", "社保政策问答、增减员材料校验、薪酬个税测算、社保异常预警、客户月报生成。"),
        ("乐接活", "灵活用工准入、四流合一检查、灵工签约助手、结算异常识别、税务路径建议。"),
        ("金蜻蜓", "骑手招聘、运单异常、保险投保、理赔材料审核、新职伤申报、发薪报税检查。"),
        ("微办公/超级 HR", "员工自助问答、审批助手、公告生成、考勤异常、合同模板生成。"),
        ("小霖数字员工", "从单点问答升级为跨系统任务助手，统一承接企业文档生成、人财法咨询、属地政策查询。"),
        ("淮安政企项目", "AI 就业平台、AI 社保智能体、数字仲裁辅助、人才画像、岗位画像、产业报告。"),
    ]
    add_matrix(doc, ["业务线", "Agent 机会"], rows, [1.3, 5.2])

    add_h(doc, "附录 B：指标体系", 1)
    rows = [
        ("效率", "平均处理时长、自动生成比例、人工节省工时、工单关闭周期。"),
        ("质量", "答案准确率、材料缺失识别率、异常召回率、误报率、人工修正率。"),
        ("风控", "高风险漏判率、准入拦截率、审计完整率、人工确认覆盖率。"),
        ("业务", "客户满意度、续费率、投诉率、销售方案采纳率、项目交付准时率。"),
        ("产品", "日活用户、调用次数、复用场景数、知识库更新频率、工具调用成功率。"),
    ]
    add_matrix(doc, ["指标类别", "建议指标"], rows, [1.2, 5.3])

    add_h(doc, "附录 C：业务脑图参考", 1)
    add_p(doc, "下图来自“数智创新数字化 2026”业务脑图，反映了集团产研中后台、淮安业务、本地拓展、技术底座、统一数据、业务系统与安全 AI 的关系。文档中的 Agent 场景设计以该业务地图为重要参考。")
    if MINDMAP.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(MINDMAP), width=Inches(6.5))
    else:
        add_p(doc, "未找到业务脑图图片。", color=MUTED)

    add_h(doc, "附录 D：建议下一步会议议程", 1)
    add_numbers(doc, [
        "确认优先 MVP：骑手理赔材料审核 Agent 或社保薪税政策助手 Agent 二选一。",
        "拉通业务负责人：金蜻蜓运营、客服交付、风控、产品、技术共同确认流程。",
        "收集历史案例：至少 50 个真实样本，覆盖正常、缺件、异常、高风险四类。",
        "定义上线边界：明确 Agent 能自动做什么、必须人工确认什么、禁止做什么。",
        "确定 90 天指标：效率、质量、风控、满意度和复用产出。",
    ])

    add_h(doc, "附录 E：骑手理赔材料审核 Agent MVP 需求样例", 1)
    add_p(doc, "如果淮安分部要快速做出一个可演示、可试用、可量化的 Agent，建议优先选择“骑手理赔材料审核 Agent”。该场景资料类型明确、人工重复工作多、风控价值明显，也容易形成业务闭环。")
    rows = [
        ("用户角色", "理赔专员、客服、运营主管、风控专员。"),
        ("触发方式", "上传案件材料、创建理赔工单、客服会话转入、批量导入待审核案件。"),
        ("输入材料", "骑手身份、订单/运单、保单、事故说明、事故照片、诊断证明、费用票据、责任认定、历史理赔记录。"),
        ("核心输出", "材料完整性结论、缺失材料清单、风险等级、建议处理路径、内部审核摘要、对外补件话术。"),
        ("自动化边界", "可自动识别和提示，不自动承诺赔付，不自动拒赔，不自动对外发送高风险结论。"),
        ("上线范围", "先选一个业务组、一个保险产品或一类高频案件试点，避免初期覆盖所有理赔类型。"),
    ]
    add_matrix(doc, ["需求项", "说明"], rows, [1.35, 5.15])

    add_h(doc, "E.1 MVP 状态流转", 2)
    rows = [
        ("待识别", "系统收到案件材料，尚未抽取结构化字段。", "OCR/解析 Agent"),
        ("待补件", "材料缺失或关键字段不清晰。", "理赔 Agent 生成补件清单，客服确认后发送。"),
        ("待复核", "存在中高风险、字段冲突或规则不确定。", "转人工专员或风控复核。"),
        ("可快审", "材料齐全、低风险、规则命中明确。", "Agent 生成摘要，人工一键确认。"),
        ("已退回", "不满足基本受理条件或重复提交。", "必须人工确认退回原因。"),
        ("已归档", "案件完成，输出和人工修改进入案例库。", "用于质检和模型评测。"),
    ]
    add_matrix(doc, ["状态", "含义", "处理动作"], rows, [1.0, 2.8, 2.7])

    add_h(doc, "E.2 输出模板", 2)
    add_bullets(doc, [
        "案件摘要：骑手姓名、事故时间、事故地点、关联运单、保单状态、主要伤情、费用金额。",
        "材料状态：已提交材料、缺失材料、疑似无效材料、需重新上传材料。",
        "规则命中：保障期间、任务期间、身份一致、票据一致、金额区间、责任类型。",
        "风险提示：身份不一致、保单失效、事故时间不匹配、重复报案、票据异常、非任务期间事故。",
        "建议动作：可快审、补件、转风控、转法务、转保险公司确认。",
        "对外话术：给骑手或客户的简洁补件说明，避免出现未经确认的赔付承诺。",
    ])

    add_h(doc, "附录 F：知识库字段与治理模板", 1)
    add_p(doc, "知识库不是把 PDF 和 Word 文件直接丢给模型。为了让 Agent 可控、可追溯、可更新，建议每条知识都至少具备以下字段。")
    rows = [
        ("知识 ID", "唯一编号，例如 POLICY_JS_HA_2026_001。", "便于引用和审计。"),
        ("标题", "知识条目的业务标题。", "例如“淮安市社保增员材料要求”。"),
        ("业务线", "薪社汇、乐接活、金蜻蜓、出海、微办公。", "支持按业务检索。"),
        ("城市/区域", "全国、省、市、区县。", "属地政策必须填写。"),
        ("适用对象", "企业、HR、骑手、灵工、客服、交付人员。", "决定回答口径。"),
        ("有效期", "生效日期、失效日期、待确认状态。", "防止过期政策误用。"),
        ("来源", "政府网站、公司 SOP、法务确认、客户合同、历史案例。", "高风险回答必须有来源。"),
        ("审核人", "业务、法务、财税或风控审核人。", "明确责任。"),
        ("风险等级", "低、中、高。", "决定是否可直接回答。"),
        ("关联工具", "可调用的系统工具或办理流程。", "支撑从回答到办事。"),
    ]
    add_matrix(doc, ["字段", "说明", "作用"], rows, [1.0, 3.0, 2.5])

    add_h(doc, "附录 G：工具接口清单样例", 1)
    add_p(doc, "Agent 要办事，必须把业务系统能力封装成稳定工具。以下是淮安金蜻蜓和薪社汇优先需要的工具接口方向。")
    rows = [
        ("get_rider_profile", "根据骑手 ID 查询身份、联系方式、入驻状态、历史案件。", "理赔、招聘、风控"),
        ("get_order_detail", "查询运单时间、地点、金额、配送状态、客户归属。", "运单异常、理赔、新职伤"),
        ("get_policy_rule", "按城市、业务线、问题类型查询政策和 SOP。", "社保薪税、客服、交付"),
        ("check_insurance_status", "查询骑手保单、保障范围、有效期和投保状态。", "保险、理赔"),
        ("create_ticket", "创建工单并写入 Agent 摘要、风险等级和待办事项。", "客服、交付、风控"),
        ("send_material_request", "生成并发送补件通知，记录发送时间和模板版本。", "理赔、社保办理"),
        ("check_four_flow", "检查合同流、业务流、资金流、票据流一致性。", "灵活用工、发薪报税"),
        ("generate_report", "按时间、项目、客户生成日报、周报、风险报告。", "管理看板"),
    ]
    add_matrix(doc, ["工具名", "能力", "适用场景"], rows, [1.5, 3.3, 1.7])

    add_h(doc, "附录 H：Prompt 与规则设计原则", 1)
    add_p(doc, "Prompt 不应承担全部业务逻辑。建议把稳定规则放进规则引擎，把可解释表达和复杂语义理解交给模型，把最终动作交给工作流。")
    add_h(doc, "H.1 系统提示词应包含的内容", 2)
    add_bullets(doc, [
        "角色：你是霖珑云科淮安分部的业务辅助 Agent，服务对象是内部客服、运营、理赔和风控人员。",
        "任务边界：你可以检索、总结、校验、生成建议，但不能替代人工做法律、税务、拒赔、劳动关系最终判断。",
        "回答要求：必须输出依据、适用范围、风险等级、下一步动作；依据不足时必须提示人工确认。",
        "数据要求：不得编造政策、不得伪造材料、不得隐藏不确定性、不得输出未经授权的个人敏感信息。",
        "转人工规则：高风险、低置信度、政策过期、字段冲突、客户投诉、法律税务争议，必须转人工。",
    ])
    add_h(doc, "H.2 规则优先级", 2)
    rows = [
        ("硬规则", "法律法规、公司红线、权限边界、禁止自动决策事项。", "必须执行，模型不能覆盖。"),
        ("业务规则", "SOP、材料清单、风控评分、办理流程、城市属地规则。", "优先于模型自由推理。"),
        ("模型判断", "语义理解、材料摘要、话术生成、复杂案例解释。", "必须给出置信度和依据。"),
        ("人工结论", "授权人员确认后的最终结果。", "作为最高优先级并反哺案例库。"),
    ]
    add_matrix(doc, ["优先级", "内容", "执行原则"], rows, [1.0, 3.1, 2.4])

    add_h(doc, "附录 I：验收清单", 1)
    rows = [
        ("功能验收", "能完成指定场景输入、识别、判断、输出、转人工、留痕全流程。"),
        ("准确性验收", "历史样本测试达到目标准确率，高风险漏判为 0 或接近 0。"),
        ("体验验收", "业务人员认为输出可用，能减少重复劳动，而不是增加审核负担。"),
        ("安全验收", "敏感信息可控，高风险动作有人审，所有关键输出可追溯。"),
        ("运维验收", "知识可更新、规则可配置、日志可查询、异常可回滚、模型版本可追踪。"),
        ("产品验收", "能沉淀成 SOP、培训材料、指标看板和后续版本路线图。"),
    ]
    add_matrix(doc, ["验收类别", "验收标准"], rows, [1.25, 5.25])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
